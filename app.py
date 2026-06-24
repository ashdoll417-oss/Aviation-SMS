from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file, current_app
from flask import send_from_directory
from flask_mail import Mail, Message
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
import os
import sys
import time
from sqlalchemy import text, extract
from sqlalchemy.exc import OperationalError, SQLAlchemyError, IntegrityError
from config import Config
from extensions import db, migrate
from models import Component, RiskAssessment, SafetyPolicy, SafetyAssurance, SafetyPromotion, EmergencyResponsePlan, HazardReport, OccurrenceReport, SafetyObjective, SafetyDrill, EmergencyDrill as ModelsEmergencyDrill, Tenant, User, LoginLog
import pandas as pd
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from datetime import datetime, timedelta, timezone

RISK_MATRIX = {
    '5A': 'Unacceptable', '5B': 'Unacceptable', '5C': 'Unacceptable', '5D': 'Tolerable', '5E': 'Acceptable',
    '4A': 'Unacceptable', '4B': 'Unacceptable', '4C': 'Tolerable', '4D': 'Acceptable', '4E': 'Acceptable',
    '3A': 'Unacceptable', '3B': 'Tolerable', '3C': 'Tolerable', '3D': 'Acceptable', '3E': 'Acceptable',
    '2A': 'Tolerable', '2B': 'Tolerable', '2C': 'Acceptable', '2D': 'Acceptable', '2E': 'Acceptable',
    '1A': 'Acceptable', '1B': 'Acceptable', '1C': 'Acceptable', '1D': 'Acceptable', '1E': 'Acceptable'
}

def create_app(config_class=Config):
    # Redirect instance path to a writable directory on Vercel (read-only filesystem elsewhere)
    if os.environ.get("VERCEL"):
        app = Flask(__name__, instance_path="/tmp")
    else:
        app = Flask(__name__)

    app.config.from_object(config_class)

    # Ensure Flask session cookie works reliably across requests.
    # Some environments can yield SECRET_KEY as None/empty, which breaks @login_required session persistence.
    if not app.config.get('SECRET_KEY'):
        app.config['SECRET_KEY'] = getattr(config_class, 'SECRET_KEY', None) or 'aviation-sms-erp-dev-secret-key-2024'

    app.secret_key = app.config['SECRET_KEY']

    # --- Email (Flask-Mail) configuration via environment variables ---
    app.config['MAIL_SERVER'] = os.environ.get('MAIL_SERVER', 'smtp.gmail.com')
    app.config['MAIL_PORT'] = 587
    app.config['MAIL_USE_TLS'] = True
    app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
    app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
    app.config['MAIL_DEFAULT_SENDER'] = os.environ.get('MAIL_USERNAME')

    mail = Mail(app)

    db.init_app(app)
    migrate.init_app(app, db)
    
    # Only run create_all + default user seeding locally.
    # On Vercel, do not touch the DB during boot to prevent 500s from transient connectivity/auth issues.
    if os.environ.get("VERCEL") is None:
        with app.app_context():
            try:
                db.create_all()
            except Exception as e:
                print(f"db.create_all() skipped/failed during startup: {e}")

        with app.app_context():
            try:
                if User.query.count() == 0:
                    default_user = User(
                        username='Admin',
                        password_hash=generate_password_hash('admin123'),
                        email='admin@aviation-sms.com',
                        role='Safety Manager'
                    )
                    db.session.add(default_user)
                    db.session.commit()
                    print("Default Admin user created")
            except Exception as e:
                print(f"Default user initialization skipped/failed during startup: {e}")

    # --- ERP / Emergency Drill Models (defined/active in app.py) ---
    # Reuse the existing table mapping from models.py to avoid duplicate-table mapping issues.

    # Production safety: verify critical tables/columns on Vercel cold starts.
    # Soft-fail: if sync fails, log the error and allow the server to continue running.
    if os.environ.get("VERCEL"):
        @app.before_request
        def _sync_schema_on_cold_start():
            if app.config.get("_schema_sync_done"):
                return
            app.config["_schema_sync_done"] = True

            try:
                import subprocess
                # Call the script as a separate process to avoid import-time side effects.
                subprocess.run(
                    [sys.executable, os.path.join(os.path.dirname(__file__), "scripts", "sync_all_tables.py")],
                    check=True,
                    capture_output=True,
                    text=True,
                )
            except Exception as e:
                app.logger.exception("Schema sync failed on cold start (soft-fail). Error: %s", e)
    class EmergencyDrill(db.Model):
        __table__ = ModelsEmergencyDrill.__table__

    def login_required(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user_id' not in session:
                flash('Please log in to access this page.')
                return redirect(url_for('login'))
            return f(*args, **kwargs)
        return decorated_function

    def _safe_get_current_user():
        """
        Best-effort user lookup based on session user_id.
        Returns None if session is missing or user cannot be fetched.
        """
        uid = session.get('user_id')
        if not uid:
            return None
        try:
            return User.query.get(uid)
        except Exception:
            return None

    def _safe_get_current_tenant():
        """
        Best-effort tenant lookup:
        - If User has a `tenant_id`, try to load Tenant model if present.
        - Otherwise return None.
        """
        user = _safe_get_current_user()
        if not user:
            return None

        tenant_id = getattr(user, 'tenant_id', None)
        if not tenant_id:
            return None

        # Tenant model may or may not exist in this codebase yet.
        tenant_cls = globals().get('Tenant') or globals().get('Tenants')
        if tenant_cls is None:
            return None

        try:
            return tenant_cls.query.get(tenant_id)
        except Exception:
            return None

    def _safe_module_enabled(module_key: str) -> bool:
        """
        Best-effort module feature flags.
        Defaults to True if tenant or flags are missing, to keep the app resilient.
        Supported module_key values:
          - 'audits' -> track_audits
          - 'risk'   -> track_risk_management
        """
        tenant = _safe_get_current_tenant()
        if tenant is None:
            return True

        if module_key == 'audits':
            return bool(getattr(tenant, 'track_audits', True))
        if module_key == 'risk':
            return bool(getattr(tenant, 'track_risk_management', True))

        # Unknown module: default to True
        return True

    def require_module(module_key: str):
        """
        Lightweight route decorator to guard module pages using best-effort
        tenant subscription flags.

        Uses inject_tenant_features() flags (via _safe_module_enabled()).
        If a module is disabled: flash + redirect to /dashboard.
        """
        def decorator(f):
            @wraps(f)
            def wrapper(*args, **kwargs):
                enabled = _safe_module_enabled(module_key)
                if not enabled:
                    flash('This module is not active for your organization\'s subscription plan.')
                    return redirect(url_for('dashboard'))
                return f(*args, **kwargs)
            return wrapper
        return decorator

    def send_reporter_feedback(hazard_id):
        hazard = HazardReport.query.get_or_404(hazard_id)
        message = f'Your report regarding {hazard.unsafe_event} has been reviewed. Action taken: {hazard.safety_actions or "None"}. Status: {hazard.status}.'
        flash(message)

    @app.route('/')
    def index():
        if 'user_id' in session:
            return redirect(url_for('dashboard'))
        return redirect(url_for('login'))

    @app.route('/login', methods=['GET', 'POST'])
    def login():
        if request.method == 'POST':
            username = request.form['username']
            password = request.form['password']

            # Retry only the DB lookup during transient serverless connectivity issues.
            attempts = 3
            delays = [0.2, 0.6]  # before retry 1 and retry 2 respectively
            last_err = None

            for attempt in range(attempts):
                try:
                    user = User.query.filter_by(username=username).first()
                    break
                except SQLAlchemyError as e:
                    # Prevent session pollution on transient failures.
                    last_err = e
                    db.session.rollback()
                    user = None
                    if attempt < attempts - 1:
                        time.sleep(delays[attempt])
                    continue

            if user and check_password_hash(user.password_hash, password):
                session['user_id'] = user.id
                login_log = LoginLog(username=username, ip_address=request.remote_addr)
                db.session.add(login_log)
                db.session.commit()
                flash('Login successful!')
                return redirect(url_for('dashboard'))

            # Do not leak connection/host details to the user.
            if last_err is not None:
                flash('Temporary database connectivity issue—please try again.')
            else:
                flash('Invalid username or password.')

        return render_template('login.html')

    @app.route('/register', methods=['GET', 'POST'])
    def register():
        if request.method == 'POST':
            form_company_name = request.form.get('company_name', '').strip()
            form_full_name = request.form.get('full_name', '').strip()
            form_email = request.form.get('email', '').strip().lower()
            form_password = request.form.get('password', '')

            if not form_company_name or not form_full_name or not form_email or not form_password:
                flash('All fields are required.')
                return redirect(url_for('register'))

            # Basic password minimum (avoid empty/too-short)
            if len(form_password) < 6:
                flash('Password must be at least 6 characters.')
                return redirect(url_for('register'))

            hashed_password = generate_password_hash(form_password)

            # Multi-tenant onboarding:
            # 1) Create tenant first
            try:
                TenantCls = Tenant

                # Optional uniqueness guard to return friendly error instead of DB exception
                existing_tenant = TenantCls.query.filter_by(company_name=form_company_name).first()
                if existing_tenant:
                    flash('An organization with this company name already exists. Please sign in.')
                    return redirect(url_for('login'))

                existing_user = User.query.filter_by(email=form_email).first()
                if existing_user:
                    flash('An account with this email already exists. Please sign in.')
                    return redirect(url_for('login'))

                new_tenant = TenantCls(
                    company_name=form_company_name,
                    track_audits=True,
                    track_risk_management=True
                )
                db.session.add(new_tenant)
                db.session.commit()

                # 2) Create admin user linked to that tenant
                # Ensure tenant_id is assigned using the generated Tenant UUID dynamically.
                # Works whether Tenant.id is a UUID object or already a string.
                tenant_id_value = str(new_tenant.id)

                # Use a deterministic username for login compatibility. Current login expects `username`.
                # We'll set username to email local-part if present, else email.
                derived_username = form_email.split('@', 1)[0] or form_email

                # Ensure derived username uniqueness
                if User.query.filter_by(username=derived_username).first():
                    derived_username = form_email

                new_user = User(
                    username=derived_username,
                    email=form_email,
                    password_hash=hashed_password,
                    tenant_id=tenant_id_value,
                    role='Administrator'
                )
                db.session.add(new_user)
                db.session.commit()

                # Safe login session
                session['user_id'] = new_user.id
                flash('Organization registered successfully!')
                return redirect(url_for('dashboard'))

            except IntegrityError:
                db.session.rollback()
                flash('This organization name is already registered.')
                return redirect(url_for('register'))

            except SQLAlchemyError:
                db.session.rollback()
                flash('Registration failed due to a temporary database issue. Please try again.')
                return redirect(url_for('register'))

        return render_template('register.html')

    @app.route('/logout')
    def logout():
        session.pop('user_id', None)
        flash('You have been logged out.')
        return redirect(url_for('login'))

    @app.route('/favicon.ico')
    def favicon_ico():
        return send_from_directory('static', 'favicon.ico', mimetype='image/x-icon')

    @app.route('/favicon.png')
    def favicon_png():
        return send_from_directory('static', 'favicon.png', mimetype='image/png')

    @app.route('/dashboard')
    @login_required
    def dashboard():
        user = User.query.get(session['user_id'])

        # Multi-tenant isolation (best-effort):
        # If we can derive a tenant_id from the logged-in user and SafetyAssurance supports tenant_id,
        # filter tenant-specific queries. For other models, we keep existing behavior unless
        # their tenant filtering is explicitly supported in the codebase.
        # Strict tenant isolation for SafetyAssurance:
        # - If tenant_id is missing/unassigned: return absolutely nothing.
        # - No global fallbacks.
        active_user = _safe_get_current_user()
        tenant_id = getattr(active_user, 'tenant_id', None)

        components = Component.query.all()
        risks = RiskAssessment.query.all()
        policies = SafetyPolicy.query.all()
        promotions = SafetyPromotion.query.all()
        erps = EmergencyResponsePlan.query.all()
        hazards = HazardReport.query.all()
        occurrences = OccurrenceReport.query.all()
        objectives = SafetyObjective.query.all()
        drills = SafetyDrill.query.all()

        # SafetyAssurance table has no tenant_id column in models.py; never filter by tenant_id.
        assurances = []

        return render_template(
            'dashboard.html',
            user=user,
            components=components,
            risks=risks,
            policies=policies,
            assurances=assurances,
            promotions=promotions,
            erps=erps,
            hazards=hazards,
            occurrences=occurrences,
            objectives=objectives,
            drills=drills
        )

    @app.route('/export/excel')
    def export_excel():
        components = Component.query.all()
        risks = RiskAssessment.query.all()
        components_df = pd.DataFrame([{'ID': c.id, 'Name': c.name, 'Serial Number': c.serial_number, 'Part Number': c.part_number, 'Install Date': c.install_date, 'Due Date': c.due_date, 'Status': c.status.value if c.status else None} for c in components])
        risks_df = pd.DataFrame([{'ID': r.id, 'Hazard Description': r.hazard_description, 'Probability': r.probability, 'Severity': r.severity, 'Risk Level': r.risk_level.value if r.risk_level else None, 'Mitigation Plan': r.mitigation_plan} for r in risks])
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            components_df.to_excel(writer, sheet_name='Components', index=False)
            risks_df.to_excel(writer, sheet_name='Risk Assessments', index=False)
        output.seek(0)
        return send_file(output, download_name='aviation_sms_exports.xlsx', as_attachment=True)

    @app.route('/export/pdf')
    def export_pdf():
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        p.setFont("Helvetica", 12)
        p.drawString(100, 750, "Aviation SMS Risk Assessment Report")
        y = 700
        risks = RiskAssessment.query.all()
        for r in risks:
            if y < 100:
                p.showPage()
                p.setFont("Helvetica", 12)
                y = 750
            p.drawString(100, y, f"ID: {r.id} - {r.hazard_description[:50]}")
            y -= 20
        p.save()
        buffer.seek(0)
        return send_file(buffer, download_name='risk_report.pdf', as_attachment=True)

    @app.route('/export/erp/pdf')
    def export_erp_pdf():
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        p.setFont("Helvetica", 12)
        p.drawString(100, 800, "Emergency Response Plans")
        y = 750
        erps = EmergencyResponsePlan.query.all()
        for erp in erps:
            if y < 100:
                p.showPage()
                p.setFont("Helvetica", 12)
                y = 800
            p.drawString(100, y, f"Plan: {erp.plan_name}")
            y -= 20
        p.save()
        buffer.seek(0)
        return send_file(buffer, download_name='erp_report.pdf', as_attachment=True)

    @app.route('/export/risk/excel')
    def export_risk_excel():
        risks = RiskAssessment.query.all()
        risk_data = []
        for r in risks:
            row = {'id': r.id, 'hazard': r.hazard_description, 'probability': r.probability, 'severity': r.severity, 'risk': r.risk_level.value if r.risk_level else '', 'mitigation': r.mitigation_plan}
            risk_data.append(row)
        df = pd.DataFrame(risk_data)
        output = BytesIO()
        df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)
        return send_file(output, download_name='risk_assessment.xlsx', as_attachment=True)

    @app.route('/erp/update', methods=['GET', 'POST'])
    @login_required
    def update_erp():
        erp = EmergencyResponsePlan.query.get_or_404(request.args.get('id'))
        if request.method == 'POST':
            erp.plan_name = request.form['plan_name']
            erp.description = request.form['description']
            erp.coordinator = request.form['coordinator']
            erp.phone = request.form['phone']
            erp.last_reviewed = datetime.utcnow()
            db.session.commit()
            flash('Emergency Response Plan updated!')
            return redirect(url_for('dashboard'))
        return render_template('erp_update.html', erp=erp)

    @app.route('/erp/new', methods=['GET', 'POST'])
    @login_required
    def new_erp():
        if request.method == 'POST':
            erp_image = request.files.get('erp_image')

            image_rel_path = None
            if erp_image and erp_image.filename != '':
                base_dir = os.path.join('static', 'uploads', 'erp')
                os.makedirs(base_dir, exist_ok=True)

                filename = secure_filename(erp_image.filename)
                # Basic uniqueness to avoid collisions
                name_root, ext = os.path.splitext(filename)
                filename = f"{name_root}_{int(datetime.utcnow().timestamp())}{ext}"

                save_path = os.path.join(base_dir, filename)
                erp_image.save(save_path)

                image_rel_path = os.path.join('uploads', 'erp', filename).replace('\\', '/')

            new_plan = EmergencyResponsePlan(
                plan_name=request.form.get('plan_name'),
                description=request.form.get('description'),
                coordinator=request.form.get('coordinator'),
                phone=request.form.get('phone')
            )

            # If the model supports an image path column, store it
            # (best-effort; won’t break if the column doesn’t exist)
            if image_rel_path is not None and hasattr(new_plan, 'erp_image'):
                setattr(new_plan, 'erp_image', image_rel_path)

            db.session.add(new_plan)
            db.session.commit()
            flash('Emergency Response Plan created!')
            return redirect(url_for('dashboard'))
        return render_template('log_erp.html')

    @app.route('/erp/list', methods=['GET', 'POST'])
    @login_required
    def list_erp():
        if request.method == 'POST':
            # Handle photo file upload if provided
            photo_file = request.files.get('erp_image')
            filename = None
            if photo_file and photo_file.filename != '':
                import os
                upload_dir = os.path.join('static', 'uploads')
                os.makedirs(upload_dir, exist_ok=True)
                filename = photo_file.filename
                photo_file.save(os.path.join(upload_dir, filename))

            # Create Drill entry matching AISL-QD-022 format
            new_drill = EmergencyDrill(
                report_ref=request.form.get('report_ref'),
                date=request.form.get('date'),
                component_type=request.form.get('component_type') or "N/A",
                time_since_new=request.form.get('time_since_new') or "N/A",
                ac_reg=request.form.get('ac_reg') or "N/A",
                time_since_oh=request.form.get('time_since_oh') or "N/A",
                part_no=request.form.get('part_no') or "N/A",
                client_name=request.form.get('client_name') or "N/A",
                serial_no=request.form.get('serial_no') or "N/A",
                description=request.form.get('description'),
                initial_findings=request.form.get('initial_findings'),
                other_info=request.form.get('other_info'),
                conclusion=request.form.get('conclusion'),
                reported_by=request.form.get('reported_by'),
                photo_path=filename
            )
            db.session.add(new_drill)
            db.session.commit()
            return redirect(url_for('list_erp'))

        drills = EmergencyDrill.query.all()
        return render_template('list_erp.html', drills=drills)

    @app.route('/erp/delete/<int:id>', methods=['POST', 'GET'])
    @login_required
    def delete_erp_entry(id):
        entry = EmergencyDrill.query.get_or_404(id)
        db.session.delete(entry)
        db.session.commit()
        return redirect(url_for('list_erp'))

    @app.route('/erp/edit/<int:id>', methods=['GET', 'POST'])
    @login_required
    def edit_erp_entry(id):
        entry = EmergencyDrill.query.get_or_404(id)
        if request.method == 'POST':
            entry.report_ref = request.form.get('report_ref')
            entry.date = request.form.get('date')
            entry.description = request.form.get('description')
            entry.initial_findings = request.form.get('initial_findings')
            entry.conclusion = request.form.get('conclusion')
            entry.reported_by = request.form.get('reported_by')
            db.session.commit()
            return redirect(url_for('list_erp'))
        return render_template('edit_drill.html', drill=entry)

    @app.route('/risk/new', methods=['GET', 'POST'])
    @login_required
    def new_risk():
        if request.method == 'POST':
            probability = int(request.form['probability'])
            severity = int(request.form['severity'])
            risk_key = f"{severity}{probability}"
            risk_level = RISK_MATRIX.get(risk_key, 'Unknown')
            new_risk = RiskAssessment(
                hazard_description=request.form['hazard_description'],
                probability=probability,
                severity=severity,
                risk_level=risk_level,
                mitigation_plan=request.form['mitigation_plan']
            )
            db.session.add(new_risk)
            db.session.commit()
            flash('Risk assessment created!')
            return redirect(url_for('dashboard'))
        return render_template('new_risk.html')

    @app.route('/risk/assessment')
    @login_required
    def risk_assessment():
        risks = RiskAssessment.query.all()
        return render_template('risk_assessment.html', risks=risks)

    @app.route('/risk/assessment/new')
    @login_required
    def risk_assessment_new():
        risks = RiskAssessment.query.all()
        return render_template('risk_assessment_new.html', risks=risks)

    @app.route('/hazard/report', methods=['GET', 'POST'])
    @login_required
    def report_hazard():
        if request.method == 'POST':
            hazard = HazardReport(
                unsafe_event=request.form['unsafe_event'],
                location=request.form['location'],
                description=request.form['description'],
                reported_by=session.get('user_id'),
                status='Reported'
            )
            db.session.add(hazard)
            db.session.commit()
            flash('Hazard reported successfully!')
            return redirect(url_for('dashboard'))
        return render_template('report_hazard.html')

    @app.route('/hazard/assess/<int:hazard_id>', methods=['GET', 'POST'])
    @login_required
    def assess_hazard(hazard_id):
        hazard = HazardReport.query.get_or_404(hazard_id)
        if request.method == 'POST':
            hazard.status = request.form['status']
            hazard.safety_actions = request.form['safety_actions']
            hazard.assessed_by = session.get('user_id')
            db.session.commit()
            send_reporter_feedback(hazard_id)
            return redirect(url_for('dashboard'))
        return render_template('assess_hazard.html', hazard=hazard)

    @app.route('/hazard/report/list')
    @login_required
    def report_hazard_list():
        hazards = HazardReport.query.all()
        return render_template('report_hazard.html', hazards=hazards)

    @app.route('/hazard/close/<int:hazard_id>', methods=['GET', 'POST'])
    @login_required
    def close_hazard(hazard_id):
        hazard = HazardReport.query.get_or_404(hazard_id)
        if request.method == 'POST':
            hazard.status = 'Closed'
            hazard.closure_date = datetime.utcnow()
            hazard.closure_comment = request.form.get('closure_comment')
            db.session.commit()
            send_reporter_feedback(hazard_id)
            return redirect(url_for('dashboard'))
        return render_template('close_hazard.html', hazard=hazard)

    @app.route('/occurrence/report', methods=['GET', 'POST'])
    @login_required
    def report_occurrence():
        if request.method == 'POST':
            occ = OccurrenceReport(
                occurrence_type=request.form['occurrence_type'],
                date_time=datetime.strptime(request.form['date_time'], '%Y-%m-%dT%H:%M'),
                location=request.form['location'],
                flight_number=request.form['flight_number'],
                description=request.form['description'],
                reported_by=session.get('user_id'),
                status='Reported'
            )
            db.session.add(occ)
            db.session.commit()
            flash('Occurrence reported!')
            return redirect(url_for('dashboard'))
        return render_template('report_occurrence.html')

    @app.route('/occurrence/form')
    @login_required
    def occurrence_form():
        return render_template('occurrence_form.html')

    @app.route('/inventory')
    @login_required
    def inventory():
        components = Component.query.all()
        return render_template('inventory.html', components=components)

    @app.route('/drills', methods=['GET', 'POST'])
    @login_required
    def manage_drills():
        if request.method == 'POST':
            new_drill = EmergencyDrill(
                report_ref=request.form.get('report_ref'),
                date=request.form.get('date'),
                component_type=request.form.get('component_type'),
                time_since_new=request.form.get('time_since_new'),
                ac_reg=request.form.get('ac_reg'),
                time_since_oh=request.form.get('time_since_oh'),
                part_no=request.form.get('part_no'),
                client_name=request.form.get('client_name'),
                serial_no=request.form.get('serial_no'),
                description=request.form.get('description'),
                initial_findings=request.form.get('initial_findings'),
                other_info=request.form.get('other_info'),
                conclusion=request.form.get('conclusion'),
                reported_by=request.form.get('reported_by')
            )
            db.session.add(new_drill)
            db.session.commit()
            flash('Fire Drill report saved successfully!')
            return redirect(url_for('manage_drills'))

        drills = EmergencyDrill.query.all()
        return render_template('drills.html', drills=drills)

    @app.route('/drills/add', methods=['GET', 'POST'])
    @login_required
    def add_drill():
        if request.method == 'POST':
            drill = SafetyDrill(
                drill_type=request.form['drill_type'],
                description=request.form['description'],
                location=request.form['location'],
                scheduled_date=datetime.strptime(request.form['scheduled_date'], '%Y-%m-%dT%H:%M'),
                conducted_by=session.get('user_id')
            )
            db.session.add(drill)
            db.session.commit()
            flash('Drill scheduled!')
            return redirect(url_for('drills'))
        return render_template('edit_drill.html')

    @app.route('/drills/edit/<int:drill_id>', methods=['GET', 'POST'])
    @login_required
    def edit_drill(drill_id):
        drill = SafetyDrill.query.get_or_404(drill_id)
        if request.method == 'POST':
            drill.drill_type = request.form['drill_type']
            drill.description = request.form['description']
            drill.location = request.form['location']
            drill.scheduled_date = datetime.strptime(request.form['scheduled_date'], '%Y-%m-%dT%H:%M')
            drill.completed = 'completed' in request.form
            db.session.commit()
            flash('Drill updated!')
            return redirect(url_for('drills'))
        return render_template('edit_drill.html', drill=drill)

    @app.route('/manage_objectives', methods=['GET', 'POST'])
    @login_required
    def manage_objectives():
        if request.method == 'POST':
            new_obj = SafetyObjective(
                customer_no=request.form.get('customer_no'),
                operator_id=request.form.get('operator_id'),
                text=request.form.get('objective_text')
            )
            db.session.add(new_obj)
            db.session.commit()
            flash('Safety Objective added successfully!')
            return redirect(url_for('manage_objectives'))
        
        objectives = SafetyObjective.query.all()
        return render_template('manage_objectives.html', objectives=objectives)

    @app.route('/delete_objective/<int:id>')
    @login_required
    def delete_objective(id):
        objective = SafetyObjective.query.get_or_404(id)
        db.session.delete(objective)
        db.session.commit()
        flash('Objective deleted.')
        return redirect(url_for('manage_objectives'))

    @app.route('/safety/assurance', methods=['GET'])
    @login_required
    @require_module('audits')
    def safety_assurance():
        from sqlalchemy import text

        try:
            db.session.rollback() # Ensure transaction state is fresh

            query = text("""
                SELECT id, audit_date, target_month, audit_scope, status, finding_details,
                       auditee_email, auditee_responder_name, auditee_remarks,
                       root_causes, immediate_corrective_action, system_alteration,
                       auditee_signature_name, auditee_signed_date,
                       next_audit_date AS next_audit
                FROM safety_assurance
                ORDER BY audit_date DESC
            """)
            records_result = db.session.execute(query).mappings().all()
            latest = records_result[0] if records_result else None
            assurances = records_result  # back-compat variable name
        except Exception as e:
            db.session.rollback()
            print(f"CRITICAL GET QUERY FAILED: {str(e)}")
            raise e

        return render_template(
            'safety_assurance.html',
            records=records_result,
            latest=latest,
            assurances=assurances
        )

    @app.route('/safety/assurance/delete/<int:audit_id>', methods=['POST'])
    @login_required
    @require_module('audits')
    def delete_safety_audit(audit_id):
        from sqlalchemy import text

        delete_sql = text("""
            DELETE FROM safety_assurance
            WHERE id = :audit_id
        """)

        result = db.session.execute(delete_sql, {"audit_id": audit_id})
        db.session.commit()

        if getattr(result, "rowcount", 0) > 0:
            flash("Audit record has been deleted successfully.", "success")
        else:
            flash("Unauthorized deletion request or record not found.", "danger")

        return redirect(url_for('safety_assurance'))

    @app.route('/safety/assurance/report/<int:audit_id>', methods=['POST'])
    @login_required
    @require_module('audits')
    def send_audit_report_email(audit_id):
        import io
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from sqlalchemy import text
        from flask import request, flash, redirect, url_for

        # 2. Extract input values safely from the modal view post submission
        finding_details = request.form.get('finding_details', '').strip()
        recipient_email = request.form.get('recipient_email', '').strip()

        if not recipient_email:
            flash("A valid recipient email address is required to process the report delivery.", "danger")
            return redirect(url_for('safety_assurance'))

        try:
            # 3. Save findings and update recipient targets directly inside safety_assurance
            update_sql = text("""
                UPDATE safety_assurance
                SET finding_details = :finding_details,
                    auditee_email = :recipient_email,
                    status = 'Open'
                WHERE id = :audit_id
            """)
            update_result = db.session.execute(update_sql, {
                "finding_details": finding_details,
                "recipient_email": recipient_email,
                "audit_id": audit_id
            })
            db.session.commit()

            # If nothing was updated, don't proceed to email/docx generation
            if getattr(update_result, "rowcount", 0) < 1:
                flash("Audit record not found or not accessible.", "danger")
                return redirect(url_for('safety_assurance'))

            # 4. Fetch updated parameters to sync into the .docx layout fields
            query = text("""
                SELECT id, audit_date, target_month, audit_scope, status,
                       auditee_responder_name, auditee_remarks, finding_details, auditee_email
                FROM safety_assurance
                WHERE id = :audit_id
            """)
            record = db.session.execute(query, {"audit_id": audit_id}).mappings().first()

            if not record:
                flash("Audit record not found after update.", "danger")
                return redirect(url_for('safety_assurance'))
            
            # 5. Initialize Document and set margins to prevent table overflowing layout limits
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(10.5)
            
            # =========================
            # AISL-SD-011 STRUCTURE
            # =========================

            # Top Header Table (2 Columns, 4 Rows)
            top_table = doc.add_table(rows=4, cols=2)
            top_table.style = 'Table Grid'

            record_audit_date = record['audit_date'] or 'N/A'
            record_audit_id = record['id']
            record_audit_ref = f"SMS/ASR/{record_audit_id}"

            record_scope = record['audit_scope'] or 'N/A'
            record_target_month = record['target_month'] or 'N/A'

            auditee_email = record['auditee_email'] or 'N/A'
            verified_responder = record.get('auditee_responder_name') or ''
            verified_responder = verified_responder.strip() if isinstance(verified_responder, str) else verified_responder
            verified_responder = verified_responder if verified_responder else "Pending Response"

            findings_present = True if (record.get('finding_details') and str(record.get('finding_details')).strip()) else False
            findings_count = 1 if findings_present else 0

            status_val = record.get('status') or 'Open'

            top_table.rows[0].cells[0].text = f"{record_audit_date}"
            top_table.rows[0].cells[1].text = f"{record_audit_ref}"

            top_table.rows[1].cells[0].text = f"{record_scope}"
            top_table.rows[1].cells[1].text = f"{record_target_month}"

            top_table.rows[2].cells[0].text = f"{auditee_email}"
            top_table.rows[2].cells[1].text = f"{verified_responder}"

            top_table.rows[3].cells[0].text = f"{findings_count}"
            top_table.rows[3].cells[1].text = f"{status_val}"

            # Exact Section: Introduction (Bold) + exact text
            h_intro = doc.add_paragraph("Introduction")
            h_intro.runs[0].bold = True
            doc.add_paragraph(
                "The Internal Safety Audit was scheduled and carried out. "
                "This audit report entails the summary of the audit carried out."
            )

            # Exact Section: Reference Documents (Bold) + include required items
            h_ref = doc.add_paragraph("Reference Documents")
            h_ref.runs[0].bold = True
            doc.add_paragraph(
                "Internal Safety Audit Checklist, Relevant Authority Regulations, "
                "Industry best practices, and Company procedures manuals "
                "(MPM: AISL-001, SMSM: AISL-005, Training Manual: AISL-006, OSH Manual: AISL-007, "
                "Workshop Procedures Manual: AISL-009)."
            )

            # Exact Section: General Comments/Observation (Bold) + exact text
            h_gen = doc.add_paragraph("General Comments/Observation")
            h_gen.runs[0].bold = True
            doc.add_paragraph(
                "The workshops are kept clean and well ventilated. "
                "There are enough resources to perform the work required and safety standards are well observed "
                "with emphasis given on the protection of personnel from hazards they are exposed to."
            )

            # Findings Text
            h_find = doc.add_paragraph("Details of Raised Non-Conformities & Actions")
            h_find.runs[0].bold = True
            doc.add_paragraph(f"{record['finding_details'] or ''}")

            # Analysis & Classification Table (Standard block)
            analysis_table = doc.add_table(rows=2, cols=5)
            analysis_table.style = 'Table Grid'

            # Header row
            analysis_headers = ["#", "Issue/s", "Intolerable", "Tolerable", "Acceptable"]
            for i, hdr in enumerate(analysis_headers):
                analysis_table.rows[0].cells[i].text = hdr

            # Row 1
            issue_text = record.get('finding_details') or ''
            analysis_table.rows[1].cells[0].text = "1."
            analysis_table.rows[1].cells[1].text = str(issue_text).strip()
            analysis_table.rows[1].cells[2].text = "NIL"
            analysis_table.rows[1].cells[3].text = "NIL"
            analysis_table.rows[1].cells[4].text = "NIL"

            # Footer legend
            doc.add_paragraph(
                "Intolerable: immediate action (7 days) | Tolerable: 30 to 60 days | "
                "Acceptable: Acceptable as is."
            )

            # Signature Matrix Footer
            sig_title = doc.add_paragraph("Authorization Matrix")
            sig_title.runs[0].bold = True
            p_sig = doc.add_paragraph()
            p_sig.add_run("Auditor Signature / Date: ___________________").bold = True
            doc.add_paragraph("Auditee Signature / Date: ___________________").runs[0].bold = True
            
            # Flush document content safely into byte stream memory segment
            target_stream = io.BytesIO()
            doc.save(target_stream)
            target_stream.seek(0)
            
            # 6. Dispatch email notification message block using global platform mail instance
            from flask_mail import Message
            mail = current_app.extensions.get('mail')
            if mail is None:
                flash("Email service is not configured. Unable to dispatch report.", "danger")
                return redirect(url_for('safety_assurance'))
            
            clean_filename = f"Audit_Report_ID_{record['id']}_{record['audit_scope']}.docx".replace(" ", "_")
            
            # Create a public response token and include it in response link
            import secrets, datetime
            token = secrets.token_urlsafe(32)
            token_expires_at = datetime.datetime.utcnow() + datetime.timedelta(days=30)

            try:
                db.session.execute(
                    text("UPDATE safety_assurance SET public_respond_token = :token, public_respond_token_expires_at = :exp WHERE id = :audit_id"),
                    {"token": token, "exp": token_expires_at, "audit_id": record['id']}
                )
                db.session.commit()
            except Exception:
                db.session.rollback()

            response_link = f"https://aviation-sms.vercel.app/public/safety/respond/{record['id']}?token={token}"


            msg = Message(
                subject=f"Action Required: Internal Safety Audit Report - {record['audit_scope']}",
                recipients=[recipient_email],
                body=(
                    "Dear Team,\n\n"
                    f"Please find attached the formal Internal Safety Audit Report regarding the recent {record['audit_scope']} evaluation.\n\n"
                    "CRITICAL ACTION REQUIRED:\n"
                    "Kindly use the secure link below to log your formal root cause analysis, immediate corrective actions, and digital signature to close out these findings:\n"
                    f"👉 {response_link}\n\n"
                    "Best Regards,\n"
                    "Head of Safety Office"
                )
            )
            
            msg.attach(
                filename=clean_filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                data=target_stream.read()
            )
            mail.send(msg)
            
            flash(f"Audit report finalized and emailed successfully to {recipient_email}!", "success")
            
        except Exception as e:
            print(f"Mailing system transaction crash context trace: {str(e)}")
            flash(f"System failed to transmit report safely via email: {str(e)}", "danger")
            
        return redirect(url_for('safety_assurance'))

    @app.route('/safety/download-plan/<int:record_id>')
    @login_required
    def download_plan(record_id: int):
        import base64

        record = SafetyAssurance.query.get_or_404(record_id)

        # Optional: enforce user-level access (keeps consistent with existing page filtering).
        user_id = session.get('user_id')
        if getattr(record, 'user_id', None) is not None and record.user_id != user_id:
            flash('You do not have access to this plan.')
            return redirect(url_for('safety_assurance'))

        if not record.audit_plan_data:
            flash('No audit plan data found for this record.')
            return redirect(url_for('safety_assurance'))

        raw_bytes = base64.b64decode(record.audit_plan_data)

        return send_file(
            BytesIO(raw_bytes),
            as_attachment=True,
            download_name=record.audit_plan_filename or f'safety_assurance_plan_{record.id}'
        )

    @app.route('/safety/assurance/send/<audit_id>', methods=['GET', 'POST'])
    @login_required
    @require_module('audits')
    def send_audit_report(audit_id):
        # Securely look up the specific audit by its unique ID (no tenant_id column exists in SafetyAssurance model)
        sql = text("SELECT * FROM safety_assurance WHERE id = :audit_id")
        audit = db.session.execute(sql, {"audit_id": audit_id}).mappings().first()

        if not audit:
            flash("Audit record not found.", "danger")
            return redirect(url_for('safety_assurance'))

        # Re-use our robust try/except mail execution block here to send the stylized corporate email layout
        try:
            auditee_email = audit.get('auditee_email')
            if not auditee_email:
                flash("Audit found, but no auditee email is set.", "warning")
                return redirect(url_for('safety_assurance'))

            audit_id_str = str(audit.get('id') or audit_id)

            base_url = request.host_url.rstrip('/')

            # Create a public response token and include it in accept/reschedule links
            import secrets, datetime
            token = secrets.token_urlsafe(32)
            token_expires_at = datetime.datetime.utcnow() + datetime.timedelta(days=30)
            try:
                db.session.execute(
                    text("UPDATE safety_assurance SET public_respond_token = :token, public_respond_token_expires_at = :exp WHERE id = :audit_id"),
                    {"token": token, "exp": token_expires_at, "audit_id": audit_id}
                )
                db.session.commit()
            except Exception:
                db.session.rollback()

            accept_url = f"{base_url}/safety/assurance/respond/{audit_id}/accept?token={token}"
            reschedule_url = f"{base_url}/safety/assurance/respond/{audit_id}/reschedule?token={token}"

            audit_scope = audit.get('audit_scope') or 'Maintenance Facilities'
            target_month = audit.get('target_month') or 'Scheduled Month'
            next_audit_date = audit.get('next_audit_date')
            next_audit_date_str = next_audit_date.strftime('%m/%d/%Y') if hasattr(next_audit_date, 'strftime') else (str(next_audit_date) if next_audit_date else 'N/A')

            notification_body = audit.get('notification_body') or 'Please review the audit schedule.'

            msg = Message(
                subject=f"Safety Audit Report Notification: {audit_scope}",
                recipients=[auditee_email],
            )

            msg.html = f"""
<html>
<body style="font-family: Arial, sans-serif; background-color: #1a1a1a; color: #ffffff; margin: 0; padding: 20px;">
    <div style="max-width: 600px; margin: 0 auto; background-color: #2b2b2b; border: 1px solid #444; padding: 15px; text-align: center;">
        <span style="font-size: 18px; font-weight: bold; color: #ffffff; letter-spacing: 1px;">INTERNAL AUDIT NOTIFICATION - AISL-SD-001</span><br>
        <span style="font-size: 13px; color: #b3b3b3;">Aero Instrument Service Limited (AISL) | Workshop ID: K/AMO/L/016</span>
    </div>
    
    <div style="max-width: 600px; margin: 0 auto; background-color: #242424; border-left: 1px solid #444; border-right: 1px solid #444; border-bottom: 1px solid #444; padding: 20px; color: #dddddd;">
        <h4 style="color: #ffc107; border-bottom: 1px solid #444; padding-bottom: 5px; margin-top: 0;">AUDIT DETAILS</h4>
        <table style="width: 100%; color: #dddddd; font-size: 14px; margin-bottom: 20px;">
            <tr><td style="padding: 5px; width: 30%; font-weight: bold; color: #aaaaaa;">Audit Area:</td><td style="padding: 5px;">{audit_scope}</td></tr>
            <tr><td style="padding: 5px; font-weight: bold; color: #aaaaaa;">Auditor:</td><td style="padding: 5px;">Head of Safety Office</td></tr>
            <tr><td style="padding: 5px; font-weight: bold; color: #aaaaaa;">Schedule:</td><td style="padding: 5px;">{target_month}</td></tr>
            <tr><td style="padding: 5px; font-weight: bold; color: #aaaaaa;">Next Audit:</td><td style="padding: 5px;">{next_audit_date_str}</td></tr>
        </table>
        
        <p style="font-size: 14px; line-height: 1.5;">
            <strong>Preparation Notice:</strong><br>
            In accordance with <strong>AISL-005 (SMM)</strong>, please be advised of the scheduled audit. Preparation should follow the <strong>AISL-SD-002C</strong> checklist.
        </p>
        
        <div style="background-color: #2d3748; border: 1px solid #4a5568; padding: 15px; text-align: center; margin-top: 25px; border-radius: 6px;">
            <span style="font-size: 15px; font-weight: bold; color: #63b3ed;">📋 AUDIT SCHEDULE ACKNOWLEDGEMENT</span>
            <p style="font-size: 13px; color: #cbd5e0; margin: 10px 0 15px 0;">Please confirm your acceptance of this audit schedule or request rescheduling:</p>
            <a href="{accept_url}" style="background-color: #0056b3; color: white; padding: 10px 22px; text-decoration: none; font-weight: bold; border-radius: 4px; margin-right: 10px; display: inline-block; font-size: 13px;">✓ Accept Audit Schedule</a>
            <a href="{reschedule_url}" style="background-color: #dc3545; color: white; padding: 10px 22px; text-decoration: none; font-weight: bold; border-radius: 4px; display: inline-block; font-size: 13px;">✗ Request Reschedule</a>
            <br><span style="font-size: 11px; color: #a0aec0; display: inline-block; margin-top: 12px;">Audit ID: {audit_id_str}</span>
        </div>
    </div>
</body>
</html>
"""

            msg.body = (
                f"{notification_body}\n\n"
                f"Audit ID: {audit_id_str}\n"
                f"Accept: {accept_url}\n"
                f"Reschedule: {reschedule_url}\n"
            )

            # Optional: attach audit plan/checklist if present (best-effort; won't break redirect)
            audit_plan_filename = audit.get('audit_plan_filename')
            audit_plan_data = audit.get('audit_plan_data')
            if audit_plan_data:
                try:
                    raw = base64.b64decode(audit_plan_data)
                    msg.attach(
                        filename=audit_plan_filename or f'safety_assurance_plan_{audit_id_str}',
                        content_type=None,
                        data=raw,
                    )
                except Exception:
                    pass

            checklist_name = audit.get('checklist_name')
            checklist_data = audit.get('checklist_data')
            if checklist_data:
                try:
                    raw = checklist_data if isinstance(checklist_data, (bytes, bytearray)) else base64.b64decode(checklist_data)
                    msg.attach(
                        filename=checklist_name or f'safety_assurance_checklist_{audit_id_str}',
                        content_type=None,
                        data=raw,
                    )
                except Exception:
                    pass

            mail.send(msg)

            flash("Audit report successfully sent to the auditee!", "success")
        except Exception as mail_err:
            current_app.logger.error(f"Failed to dispatch audit report mail: {str(mail_err)}")
            flash("Database updated, but there was a temporary notification delivery issue.", "warning")

        return redirect(url_for('safety_assurance'))

    @app.route('/public/safety/respond/<int:audit_id>', methods=['GET', 'POST'])
    def public_audit_response_portal(audit_id):
        from sqlalchemy import text
        from flask import request, render_template, flash, redirect, url_for
        import datetime
        import hmac

        token = request.args.get('token') or request.form.get('token') or ''
        if not token:
            return "Missing token.", 403

        row = db.session.execute(
            text("SELECT id, public_respond_token, public_respond_token_expires_at FROM safety_assurance WHERE id = :audit_id"),
            {"audit_id": audit_id}
        ).mappings().first()

        if not row:
            return "Audit record link invalid or expired.", 404

        expires = row.get('public_respond_token_expires_at')
        if expires and isinstance(expires, datetime.datetime) and expires < datetime.datetime.utcnow():
            return "Token expired.", 403

        stored = row.get('public_respond_token') or ''
        if not hmac.compare_digest(stored, token):
            return "Invalid token.", 403

        if request.method == 'POST':
            conformance = request.form.get('description_of_conformance', '').strip()
            root_causes = request.form.get('root_causes', '').strip()
            corr_action = request.form.get('immediate_corrective_action', '').strip()
            sys_alteration = request.form.get('system_alteration', '').strip()
            sig_name = request.form.get('auditee_signature_name', '').strip()
            sig_date_in = request.form.get('auditee_signed_date', '').strip()

            # Allow DATE/VARCHAR column types; store as ISO date if possible, else store provided string.
            if sig_date_in:
                sig_date = sig_date_in
            else:
                sig_date = datetime.date.today().isoformat()

            try:
                update_sql = text("""
                    UPDATE safety_assurance 
                    SET description_of_conformance = :conformance,
                        root_causes = :root_causes,
                        immediate_corrective_action = :corr_action,
                        system_alteration = :sys_alteration,
                        auditee_signature_name = :sig_name,
                        auditee_signed_date = :sig_date,
                        auditee_responder_name = :sig_name,
                        auditee_remarks = :corr_action,
                        status = 'Closed'
                    WHERE id = :audit_id
                """)
                db.session.execute(update_sql, {
                    "conformance": conformance,
                    "root_causes": root_causes,
                    "corr_action": corr_action,
                    "sys_alteration": sys_alteration,
                    "sig_name": sig_name,
                    "sig_date": sig_date,
                    "audit_id": audit_id
                })
                db.session.commit()
                return render_template('public_audit_success.html')

            except Exception as e:
                db.session.rollback()
                flash(f"Error submitting corrective actions: {str(e)}", "danger")

        # GET Request: Fetch info for the public landing view frame
        query = text("SELECT id, audit_scope, audit_date, finding_details FROM safety_assurance WHERE id = :audit_id")
        record = db.session.execute(query, {"audit_id": audit_id}).mappings().first()
        if not record:
            return "Audit record link invalid or expired.", 404

        # Ensure template receives token for hidden field on POST
        return render_template('public_audit_response.html', record=record, token=token)

    @app.route('/safety/assurance/respond/<audit_id>/<action>', methods=['GET', 'POST'])
    def respond_to_audit_schedule(audit_id, action):
        from sqlalchemy import text
        import hmac
        import datetime

        token = request.args.get('token') or request.form.get('token') or ''
        if not token:
            return "Missing token.", 403

        audit = db.session.execute(
            text("SELECT * , public_respond_token, public_respond_token_expires_at FROM safety_assurance WHERE id = :audit_id"),
            {"audit_id": audit_id}
        ).mappings().first()

        if not audit:
            return "<h3>Invalid or expired audit invitation link.</h3>", 404

        expires = audit.get('public_respond_token_expires_at')
        if expires and hasattr(expires, 'tzinfo') and expires < datetime.datetime.utcnow():
            return "<h3>Token expired.</h3>", 403

        stored = audit.get('public_respond_token') or ''
        if not hmac.compare_digest(stored, token):
            return "<h3>Invalid token.</h3>", 403

        if request.method == 'POST':
            responder_name = request.form.get('responder_name')
            remarks = request.form.get('remarks')
            alt_date = request.form.get('alternative_date') or None

            # Determine status based on action click
            new_status = 'Scheduled' if action == 'accept' else 'Reschedule Requested'

            update_sql = text("""
                UPDATE safety_assurance 
                SET status = :status,
                    auditee_responder_name = :name,
                    auditee_remarks = :remarks,
                    proposed_alternative_date = :alt_date
                WHERE id = :audit_id
            """)
            db.session.execute(update_sql, {
                "status": new_status,
                "name": responder_name,
                "remarks": remarks,
                "alt_date": alt_date,
                "audit_id": audit_id
            })
            db.session.commit()

            return render_template('public_audit_success.html', audit=audit, action=action)

        # GET request shows the input form (template needs token hidden field)
        return render_template('public_audit_respond.html', audit=audit, action=action, token=token)

    @app.route('/safety/assurance', methods=['POST'])
    @login_required
    @require_module('audits')
    def safety_assurance_post():
        user_id = session.get('user_id')
        if not user_id:
            flash('Please log in to access this page.')
            return redirect(url_for('login'))

        # File helpers
        file = request.files.get('audit_plan')
        checklist_file = request.files.get('audit_checklist')

        import base64
        from datetime import datetime

        # Template uses audit_date_ui/next_audit_date_ui; keep backward-compatible fallbacks.
        audit_date_in = request.form.get('audit_date_ui') or request.form.get('audit_date')
        next_audit_date_in = request.form.get('next_audit_date_ui') or request.form.get('next_audit_date')

        # 1. Robust date parsing (handles YYYY-MM-DD and ISO-like strings with 'T')
        def _parse_ui_date(val):
            if not val:
                return None
            s = str(val).strip()
            if not s:
                return None

            # Accept ISO-like: "YYYY-MM-DDTHH:MM" / "YYYY-MM-DDTHH:MM:SS"
            if 'T' in s:
                s = s.split('T', 1)[0].strip()

            # Primary format: YYYY-MM-DD
            try:
                return datetime.strptime(s, '%Y-%m-%d').date()
            except ValueError:
                pass

            # Secondary fallback: allow YYYY/MM/DD
            try:
                return datetime.strptime(s.replace('/', '-'), '%Y-%m-%d').date()
            except Exception:
                return None

        parsed_audit_date = _parse_ui_date(audit_date_in)
        parsed_next_date = _parse_ui_date(next_audit_date_in)

        # Ensure audit_date always has a value for DB lookup/saving.
        if parsed_audit_date is None:
            parsed_audit_date = datetime.utcnow().date()

        # Set SQLAlchemy-searchable variables
        audit_date = parsed_audit_date
        next_audit_date = parsed_next_date

        status = request.form.get('status') or 'Open'
        finding_details = request.form.get('finding_details')
        auditee_email = request.form.get('auditee_email')
        notification_body = request.form.get('notification_body')
        audit_scope = request.form.get('audit_scope')
        target_month = request.form.get('target_month')
        dept_notified = True if request.form.get('department_notified') else False

        audit_plan_filename = None
        audit_plan_data = None
        if file and file.filename != '':
            allowed_ext = {'.pdf', '.docx', '.xlsx'}
            safe_name = secure_filename(file.filename)
            ext = os.path.splitext(safe_name)[1].lower()
            if ext not in allowed_ext:
                flash('Invalid file type. Allowed: PDF, DOCX, XLSX.')
                return redirect(url_for('safety_assurance'))

            file_bytes = file.read()
            file.seek(0)
            audit_plan_filename = file.filename
            audit_plan_data = base64.b64encode(file_bytes).decode('utf-8')

        checklist_name = None
        checklist_data = None
        if checklist_file and checklist_file.filename:
            checklist_name = checklist_file.filename
            checklist_data = checklist_file.read()
            checklist_file.seek(0)

        assurance = SafetyAssurance.query.filter_by(user_id=user_id, audit_date=audit_date).first()
        if assurance is None:
            assurance = SafetyAssurance(
                audit_date=audit_date,
                finding_details=finding_details,
                status=status,
                next_audit_date=next_audit_date,
                audit_scope=audit_scope,
                target_month=target_month,
                department_notified=dept_notified,
                auditee_email=auditee_email,
                notification_body=notification_body,
                checklist_name=checklist_name,
                checklist_data=checklist_data,
                user_id=user_id,
            )

        # Assign fields, including tenant-safe assignment, then commit BEFORE email code
        assurance.audit_date = parsed_audit_date
        assurance.next_audit_date = parsed_next_date

        # Tenant is intentionally NOT used for SafetyAssurance (no tenant_id column in model)

        assurance.finding_details = finding_details
        assurance.status = status
        assurance.audit_scope = audit_scope
        assurance.target_month = target_month
        assurance.department_notified = dept_notified
        assurance.auditee_email = auditee_email
        assurance.notification_body = notification_body

        if checklist_name is not None and checklist_data is not None:
            assurance.checklist_name = checklist_name
            assurance.checklist_data = checklist_data

        if audit_plan_filename is not None and audit_plan_data is not None:
            assurance.audit_plan_filename = audit_plan_filename
            assurance.audit_plan_data = audit_plan_data

        # 2. DATE FIX & DATABASE SAVE FIRST (commit immediately BEFORE executing the email code)
        try:
            # Tenant is intentionally NOT used for SafetyAssurance (no tenant_id column in model)
            assurance.status = status
            assurance.audit_scope = audit_scope
            assurance.target_month = target_month
            assurance.auditee_email = auditee_email

            db.session.add(assurance)
            db.session.commit()

            audit_id_str = str(assurance.id)
        except Exception as e:
            db.session.rollback()
            print(f"DATABASE SAVE CRASHED: {str(e)}")
            flash('Database error saving record.')
            return redirect(url_for('safety_assurance'))

        # 3. HTML EMAIL LAYOUT WITH INTERACTIVE ACTIONS + 4. TOTAL ISOLATION SAFEGUARD
        try:
            if not auditee_email:
                # No recipient => skip mailing but keep saving record
                pass
            else:
                msg = Message(
                    subject=f"New Safety Audit Notification: {request.form.get('audit_scope', 'Schedules')}",
                    recipients=[auditee_email],
                )

                record_id = audit_id_str
                base_url = request.host_url.rstrip('/')

                # Create a public response token for this audit and include it in the acceptance/reschedule links
                import secrets, datetime
                token = secrets.token_urlsafe(32)
                token_expires_at = datetime.datetime.utcnow() + datetime.timedelta(days=30)

                try:
                    db.session.execute(
                        text("UPDATE safety_assurance SET public_respond_token = :token, public_respond_token_expires_at = :exp WHERE id = :audit_id"),
                        {"token": token, "exp": token_expires_at, "audit_id": record_id}
                    )
                    db.session.commit()
                except Exception:
                    db.session.rollback()

                accept_url = f"{base_url}/safety/assurance/respond/{record_id}/accept?token={token}"
                reschedule_url = f"{base_url}/safety/assurance/respond/{record_id}/reschedule?token={token}"

                msg.html = f"""
<html>
<body style="font-family: Arial, sans-serif; background-color: #1a1a1a; color: #ffffff; margin: 0; padding: 20px;">
    <div style="max-width: 600px; margin: 0 auto; background-color: #2b2b2b; border: 1px solid #444; padding: 15px; text-align: center;">
        <span style="font-size: 18px; font-weight: bold; color: #ffffff; letter-spacing: 1px;">INTERNAL AUDIT NOTIFICATION - AISL-SD-001</span><br>
        <span style="font-size: 13px; color: #b3b3b3;">Aero Instrument Service Limited (AISL) | Workshop ID: K/AMO/L/016</span>
    </div>
    
    <div style="max-width: 600px; margin: 0 auto; background-color: #242424; border-left: 1px solid #444; border-right: 1px solid #444; border-bottom: 1px solid #444; padding: 20px; color: #dddddd;">
        <h4 style="color: #ffc107; border-bottom: 1px solid #444; padding-bottom: 5px; margin-top: 0;">AUDIT DETAILS</h4>
        <table style="width: 100%; color: #dddddd; font-size: 14px; margin-bottom: 20px;">
            <tr><td style="padding: 5px; width: 30%; font-weight: bold; color: #aaaaaa;">Audit Area:</td><td style="padding: 5px;">{assurance.audit_scope if assurance.audit_scope else 'Maintenance Facilities'}</td></tr>
            <tr><td style="padding: 5px; font-weight: bold; color: #aaaaaa;">Auditor:</td><td style="padding: 5px;">Head of Safety Office</td></tr>
            <tr><td style="padding: 5px; font-weight: bold; color: #aaaaaa;">Schedule:</td><td style="padding: 5px;">{assurance.target_month if assurance.target_month else 'Scheduled Month'}</td></tr>
        </table>
        
        <p style="font-size: 14px; line-height: 1.5;">
            <strong>Preparation Notice:</strong><br>
            In accordance with <strong>AISL-005 (SMM)</strong>, please be advised of the scheduled audit. Preparation should follow the <strong>AISL-SD-002C</strong> checklist.
        </p>
        
        <div style="background-color: #2d3748; border: 1px solid #4a5568; padding: 15px; text-align: center; margin-top: 25px; border-radius: 6px;">
            <span style="font-size: 15px; font-weight: bold; color: #63b3ed;">📋 AUDIT SCHEDULE ACKNOWLEDGEMENT</span>
            <p style="font-size: 13px; color: #cbd5e0; margin: 10px 0 15px 0;">Please confirm your acceptance of this audit schedule or request rescheduling:</p>
            <a href="{accept_url}" style="background-color: #0056b3; color: white; padding: 10px 22px; text-decoration: none; font-weight: bold; border-radius: 4px; margin-right: 10px; display: inline-block; font-size: 13px;">✓ Accept Audit Schedule</a>
            <a href="{reschedule_url}" style="background-color: #dc3545; color: white; padding: 10px 22px; text-decoration: none; font-weight: bold; border-radius: 4px; display: inline-block; font-size: 13px;">✗ Request Reschedule</a>
            <br><span style="font-size: 11px; color: #a0aec0; display: inline-block; margin-top: 12px;">Audit ID: {audit_id_str}</span>
        </div>
    </div>
</body>
</html>
"""

                msg.body = (
                    f"{notification_body or 'Please review the audit schedule.'}\n\n"
                    f"Audit ID: {audit_id_str}\n"
                    f"Accept: {accept_url}\n"
                    f"Reschedule: {reschedule_url}\n"
                )

                # Best-effort attachments (will not break redirect)
                audit_plan_file = request.files.get('audit_plan')
                audit_checklist_file = request.files.get('audit_checklist')

                if audit_plan_file and audit_plan_file.filename != '':
                    audit_plan_file.seek(0)
                    msg.attach(
                        filename=audit_plan_file.filename,
                        content_type=getattr(audit_plan_file, 'content_type', None),
                        data=audit_plan_file.read(),
                    )
                    audit_plan_file.seek(0)

                if audit_checklist_file and audit_checklist_file.filename != '':
                    audit_checklist_file.seek(0)
                    msg.attach(
                        filename=audit_checklist_file.filename,
                        content_type=getattr(audit_checklist_file, 'content_type', None),
                        data=audit_checklist_file.read(),
                    )
                    audit_checklist_file.seek(0)

                mail.send(msg)

        except Exception as mail_err:
            current_app.logger.error(f"Mail dispatch error: {mail_err}")
            flash('Safety Assurance record saved successfully.')
            return redirect(url_for('safety_assurance'))

        flash('Safety Assurance record saved successfully.')
        return redirect(url_for('safety_assurance'))

    @app.context_processor
    def inject_tenant_features():
        """
        Template context helper for feature flags.
        Defensive lookup:
          - Uses session.get('user_id') -> User record -> linked tenant fields (if present)
          - Defaults to True when user/tenant/attributes are missing to keep app resilient.
        """
        tenant_features = {
            'track_audits': True,
            'track_risk_management': True,
        }

        user = _safe_get_current_user()
        if user is None:
            return tenant_features

        # Tenant model may not exist yet; attempt best-effort loading via globals.
        tenant_id = getattr(user, 'tenant_id', None)
        if not tenant_id:
            return tenant_features

        tenant_cls = globals().get('Tenant') or globals().get('Tenants')
        if tenant_cls is None:
            return tenant_features

        try:
            tenant = tenant_cls.query.get(tenant_id)
        except Exception:
            tenant = None

        if tenant is None:
            return tenant_features

        tenant_features['track_audits'] = bool(getattr(tenant, 'track_audits', True))
        tenant_features['track_risk_management'] = bool(getattr(tenant, 'track_risk_management', True))
        return tenant_features

    return app  # This MUST be the last line of the create_app function

if __name__ == "__main__":
    app = create_app()

    # Secure startup: never enable Flask debug/reloader on Vercel by default.
    debug_env = os.environ.get("DEBUG")
    debug = (debug_env is not None and debug_env.strip().lower() in ("1", "true", "yes", "y", "on"))
    if os.environ.get("VERCEL"):
        debug = False

    port = int(os.environ.get("PORT", "5000"))
    app.run(debug=debug, port=port, use_reloader=debug)
